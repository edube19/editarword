from array import array
from audioop import add
from calendar import c
from cgi import print_arguments
from re import T, sub
from turtle import st
from docx import Document
import aspose.words as aw
#from fitz import *
from email.mime import image
from importlib.resources import path
from PIL import Image
from pyparsing import condition_as_parse_action 
from pytesseract import pytesseract 
from pdf2image import convert_from_path
from requests import delete
from recurso_prueba import*
import docx
from docx import Document
from docx.shared import Cm
from docx.shared import Pt
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import re
from unicodedata import normalize

def principalv3_prueba(ruta,string,ruta_guardar):#solo para pruebas
    #para que almacene el documento en una variable
    try:        
        doc=docx.Document(ruta) 
        print('el documento es .docx')
    except Exception :#espera el error de tipo de extension (.doc)
        nruta=corregir_extension(ruta)#cambiamos la extension a docx
        editar_extension(nruta)#eliminamos los cambios por el cambio de extension
        doc=docx.Document(nruta)#ya esta listo para trabajar
        ruta=nruta
        print('el documento es .doc')

    #eliminando los espacios en blanco doble en los documentos
    
    #leerdoc(doc)
    #if (buscar_palabra(string,doc)):#buscar si el nombre del cliente se encuentra en el documento
    if  (buscar_nombre_separado(string,doc)):
        #elimina los espacios dobles en todo el documento, asi como espacios de mas
        eliminar_doble_espacios(doc,ruta)

        #U+200E es un caracter vacio, solo para "borrar" la palabra SEÑOR NOTARIO
        editar_linea(doc,'SEÑOR NOTARIO:','‎',ruta_guardar)

        #poner en literal la cantidad de porcentaje
        porcentaje_decimalv2(doc,ruta_guardar)

        #cambiar las comas por puntos
        cambiarcomas(doc,ruta_guardar)

        #en caso haya cantidades de mas de millon
        camrbiarmillon(doc,ruta_guardar)

        print('Documento editado')
    else:
        print('El cliente no corresponde al documento o se ingreso incorrectamente')

def principalv3(ruta,string,ruta_guardar):
    #para que almacene el documento en una variable
    try:        
        doc=docx.Document(ruta) 

    except Exception :#espera el error de tipo de extension (.doc)
        nruta=corregir_extension(ruta)#cambiamos la extension a docx
        editar_extension(nruta)#eliminamos los cambios por el cambio de extension
        doc=docx.Document(nruta)#ya esta listo para trabajar
        ruta=nruta

    #eliminando los espacios en blanco doble en los documentos
    

    if (buscar_palabra(string,doc)):#buscar si el nombre del cliente se encuentra en el documento
        #elimina los espacios dobles en todo el documento, asi como espacios de mas
        eliminar_doble_espacios(doc,ruta)

        #U+200E es un caracter vacio, solo para "borrar" la palabra SEÑOR NOTARIO
        editar_linea(doc,'SEÑOR NOTARIO:','‎',ruta_guardar)

        #poner en literal la cantidad de porcentaje
        porcentaje_decimalv2(doc,ruta_guardar)

        #cambiar las comas por puntos
        cambiarcomas(doc,ruta_guardar)

        #en caso haya cantidades de mas de millon
        camrbiarmillon(doc,ruta_guardar)

        print('Documento editado')
    else:
        print('El cliente no corresponde al documento o se ingreso incorrectamente')

def cantidad_lineas(doc):
    l=len(doc.paragraphs)
    return l

def corregir_extension(ruta):#la ruta termina en .doc
    doc = aw.Document(ruta)
    b=ruta.split('.')
    ruta_docx=b.pop(0)
    ruta_docx=ruta_docx+'.docx'
    doc.save(ruta_docx)
    #print(ruta_docx)
    return ruta_docx
    
def editar_extension(ruta):
    doc=docx.Document(ruta)
    l=cantidad_lineas(doc)
    #eliminamos las lineas que salen al cambiar de doc → docx 
    doc.paragraphs[0].text=''
    doc.paragraphs[l-1].text=''
    editar_pie_pagina(doc)
    doc.save(ruta)

def editar_pie_pagina(doc):
    section = doc.sections[0] #solo hay una seccion
    footer = section.footer 
    footer.paragraphs[0].text=''

def reconstruir_string(string):
    string=string.strip()#elimina los espacios en blanco del comienzo y final
    ns = " ".join(string.split())
    return ns #retorna un string sin espacios dobles

def quitar_acentos(string):#quitar solo acentos de un string
    # Pingüino: Málãgà ês uñ̺ã cíudãd fantástica y èn Logroño me pica el... moñǫ̝̘̦̞̟̩̐̏̋͌́ͬ̚͡õ̪͓͍̦̓ơ̤̺̬̯͂̌͐͐͟o͎͈̳̠̼̫͂̊

    # Pinguino: Malaga es una ciudad fantastica y en Logroño me pica el... moñoooo
    string = re.sub(
        r"([^n\u0300-\u036f]|n(?!\u0303(?![\u0300-\u036f])))[\u0300-\u036f]+", r"\1", 
        normalize( "NFD", string), 0, re.I
    )
    string = normalize( 'NFC', string)
    #print(string)
    return string

def tasa_fija_negrita(doc,nuevostring):#en construccion
    # nuevostring es el literal del porcentaje ej->(ocho punto treinta y  siete por ciento)
    l=cantidad_lineas(doc)
    a=-1
    string='Tasa Fija'
    #string = reconstruir_string(string)
    for i in range(l):
        cadena=doc.paragraphs[i].text
        r=cadena.find(string)
        if r!=a:
            try:
                nuevacadena=cadena.split(nuevostring)
                nuevacadena[0]=''
                #doc.paragraphs[i].text=''
                nuevacadena[0]=doc.add_paragraph()
                run=doc.add_run(nuevostring)
                font = run.font
                font.name = 'Arial Narrow'
                font.size = Pt(12)
                font.bold = True #ponerlo en negrita
            except Exception as e:
                print(e)

def leerdoc(doc):
    l=cantidad_lineas(doc)
    for i in range(l):
        cadena=doc.paragraphs[i].text
        print('Parrafo N '+str(i+1))
        print(cadena)

def eliminar_linea(doc,string,ruta_guardar):#en construccion
    l=cantidad_lineas(doc)
    a=-1
    c=0
    for i in range(l):
        cadena=doc.paragraphs[i].text
        r=cadena.find(string)
        if r!=a:
            del doc.paragraphs[i]
            c=c+1
            doc.save(ruta_guardar)
            break
    if (c==0):
        print('No se encontro')

def eliminar_doble_espacios(doc,ruta):
    l=cantidad_lineas(doc)
    for i in range(l):
        cadena=doc.paragraphs[i].text
        doc.paragraphs[i].text=reconstruir_string(cadena)#elimina espacios dobles en blanco
    try:
        doc.save(ruta)
    except PermissionError:
        pass

def porcentaje_decimalv2(doc,ruta_guardar):#version final
    l=cantidad_lineas(doc)
    a=-1
    for i in range(l):
        cadena=doc.paragraphs[i].text
        if (evitar_sobreescritura(cadena,'por ciento')):
            r=cadena.find('%')
            if (r!=a):
                guardar_cadena=cadena.split('%')
                for j in range(len(guardar_cadena)-1):
                    numero=''
                    contador=1
                    cond=True
                    subcadena=guardar_cadena[j]
                    l=len(subcadena)
                    while(cond):
                        valor=subcadena[l-contador]#obtenner el valor
                        cond=valor.isdigit()#comprobar si es digito o no, si es falso acaba la ejecucion
                        if (cond):
                            numero=valor+numero#guardando el numero
                            contador=contador+1#avanza para el siguiente valor
                        elif(valor=='.'):# para casos → ab.cd%
                            contador=contador+1
                            numero='.'+numero
                            cond=True
                        elif (valor==' ' and numero==''):# para casos → ab.cd %
                            contador=contador+1
                            cond=True
                    try:
                        punto=numero.find('.')
                        if punto!=a:
                            numero_separado=numero.split('.')#separa el string segun el caracter separador
                            long_decimal=len(numero_separado[1])
                            cero='0' 
                            decimal_cero=cero*long_decimal
                            if (decimal_cero!=numero_separado[1]):
                                numeroconv1=int(numero_separado[0])#parte entera
                                numeroconv2=int(numero_separado[1])#parte decimal
                                porcentaje1=numero_to_letras(numeroconv1)
                                porcentaje2=numero_to_letras(numeroconv2)
                                porcentajeminus1=porcentaje1.lower()
                                porcentajeminus2=porcentaje2.lower()
                                nuevostring='% ('+porcentajeminus1+' punto '+porcentajeminus2+' por ciento)'
                                guardar_cadena[j]=guardar_cadena[j]+nuevostring
                            else:
                                numeroconv1=int(numero_separado[0])#parte entera
                                porcentaje1=numero_to_letras(numeroconv1)
                                porcentajeminus1=porcentaje1.lower()
                                nuevostring='% ('+porcentajeminus1+' por ciento)'
                                guardar_cadena[j]=guardar_cadena[j]+nuevostring
                            #doc.paragraphs[i].text=nueva_cadena
                        else:
                            numeroconv1=int(numero)
                            porcentaje1=numero_to_letras(numeroconv1)
                            porcentajeminus1=porcentaje1.lower()
                            nuevostring='% ('+porcentajeminus1+' por ciento)'
                            guardar_cadena[j]=guardar_cadena[j]+nuevostring
                            #doc.paragraphs[i].text=nueva_cadena
                    except Exception as e:
                        print(e)
                #cambiar formato        
                nueva_linea=''
                for parte in guardar_cadena:
                    nueva_linea=nueva_linea+parte
                #
                nueva_linea=nueva_linea.strip()
                nueva_linea=" ".join(nueva_linea.split())
                try:
                    doc.paragraphs[i].text=''
                    doc.paragraphs[i]=doc.add_paragraph()
                    run=doc.paragraphs[i].add_run(nueva_linea)
                    font = run.font
                    font.name = 'Arial Narrow'
                    font.size = Pt(12)
                except Exception as e:
                    print(e)
                #
    porcentaje_tablas(doc) 
    doc.save(ruta_guardar) 

def porcentaje_tablas(doc):
    leertabla=doc.tables
    a=-1
    for x in range(len(leertabla)):
        tabla=leertabla[x]#obtener la primera tabla
        for i in range(0,len(tabla.rows)):#filas
            for j in range(0,len(tabla.columns)):#columnas
                cadena=tabla.cell(i,j).text
                #nc=cadena[20:30]
                #print(nc)
                if (evitar_sobreescritura(cadena,'por ciento')):
                    r=cadena.find('%')
                    if (r!=a):
                #print('Se encontro en la linea '+str(i+1))
                #print(doc.paragraphs[i].text)
                        num1=cadena[r-1]
                        cond1=num1.isdigit()
                        numero=num1
                        if (cond1):
                            num2=cadena[r-2]
                            cond2=num2.isdigit()
                            if(cond2):
                                numero=num2+num1
                                num3=cadena[r-3]
                                cond3=num3.isdigit()
                                if (cond3):
                                    numero=num3+num2+num1
                        #print(numero)   
                        try:      
                            numeroconv=int(numero)
                            #intnum=int(numero)
                            porcentaje=numero_to_letras(numeroconv)
                            #porcentaje=numero_to_letras(numero)
                            porcentajeminus=porcentaje.lower()
                            nuevostring='%('+porcentajeminus+' por ciento)'
                            nueva_cadena=cadena.replace('%',nuevostring)
                            tabla.cell(i,j).text=nueva_cadena
                        except Exception as e:
                            print(e)

def editar_linea(doc,string,linea_cambiar,ruta_guardar):
    l=cantidad_lineas(doc)
    a=-1
    c=0
    for i in range(l):
        cadena=doc.paragraphs[i].text
        r=cadena.find(string)
        if r!=a:
            nlinea=cadena.replace(string,linea_cambiar)
            doc.paragraphs[i].text=nlinea
            c=c+1
    if (c==0):
        string=string.lower()
        r=cadena.find(string)
        if r!=a:
            nlinea=cadena.replace(string,linea_cambiar)
            doc.paragraphs[i].text=nlinea
            c=c+1
        else:
            print('No se encontro')
    else:
        doc.save(ruta_guardar)

def evitar_sobreescritura(cadena,string):#ponerlo dentro de un if
    cond=True
    a=-1
    r=cadena.find(string)#string = 'por ciento'
    if r!=a:
        cond=False
    return cond

def buscar_palabra(string,doc):#debe devolver un booleano
    l=cantidad_lineas(doc)
    a=-1
    c=0
    valor=False
    #string=string.upper() siemrpe sera mayuscula porq lo saca del trazer
    string = reconstruir_string(string)# elimina espacios en blanco repetidos
    print('Buscando a '+string+' en el documento')
    #stringmin=string.lower()
    for i in range(l):
        cadena=doc.paragraphs[i].text
        cadena=cadena.upper()
        cadena=quitar_acentos(cadena)
        r=cadena.find(string)
        if (r!=a):
            c=c+1
            valor=True
            break
        
    if (c==0):
        valor=buscar_en_tabla(string,doc)
        #valor=buscar_nombre_separado(string,doc)
    return valor

def buscar_nombre_separado(string,doc):
    nombre_reconstruido=string.split(' ')
    l=len(nombre_reconstruido)
    c=0
    valor=False
    while (c!=l and valor==False):
        nombre_re=' '
        for i in range(0,l-c):
            nombre_re=nombre_re+nombre_reconstruido[i]
            nombre_re=nombre_re+' '
        c+=1
        nombre_re=reconstruir_string(nombre_re)
        valor=buscar_palabra(nombre_re,doc)
    return valor

def cambiarcomas(doc,ruta_guardar):
    l=cantidad_lineas(doc)
    a=-1
    caracteres=['S/','US$']
    for i in range(l):
        cadena=doc.paragraphs[i].text
        for caracter in caracteres:
            r=cadena.find(caracter)#posicion en que se encuentra la coma
            if r!=a:
                nueva_cadena=cadena.replace(',','.')
                #doc.paragraphs[i].text=nueva_cadena
                try:
                    doc.paragraphs[i].text=''
                    doc.paragraphs[i]=doc.add_paragraph()
                    run=doc.paragraphs[i].add_run(nueva_cadena)
                    font = run.font
                    font.name = 'Arial Narrow'
                    font.size = Pt(12)
                except Exception as e:
                    print(e)
    comas_en_tabla(doc)
    
    doc.save(ruta_guardar)

def camrbiarmillon(doc,ruta_guardar):
    l=cantidad_lineas(doc)
    a=-1
    caracteres=['S/','US$']
    for i in range(l):
        cadena=doc.paragraphs[i].text
        for caracter in caracteres:
            r=cadena.find(caracter)#posicion en que se encuentra la coma
            if r!=a:
                #print('Se encontro en la linea '+str(i+1))
                #print(doc.paragraphs[i].text)
                nueva_cadena=cadena.replace('´','.')
                #print(nueva_cadena)
                doc.paragraphs[i].text=nueva_cadena
    comas_en_tabla(doc)
    doc.save(ruta_guardar)

def buscar_en_tabla(string,doc):
    leertabla=doc.tables
    resultado=0
    a=-1
    cond=False
    for x in range(len(leertabla)):
        tabla=leertabla[x]#obtener la primera tabla
        for i in range(0,len(tabla.rows)):#filas
            for j in range(0,len(tabla.columns)):#columnas
                #print(tabla.cell(i,j).text)
                cadena=tabla.cell(i,j).text
                r=cadena.find(string)
                if (r!=a):
                    print('Se encontro')
                    resultado=resultado+1
                    cond=True
                    break
    if (resultado==0):
        return cond

def comas_en_tabla(doc):
    leertabla=doc.tables
    a=-1
    caracteres=['S/','US$']
    for x in range(len(leertabla)):
        tabla=leertabla[x]#obtener la primera tabla
        for i in range(0,len(tabla.rows)):#filas
            for j in range(0,len(tabla.columns)):#columnas
                #print(tabla.cell(i,j).text)
                cadena=tabla.cell(i,j).text
                #nc=cadena[20:30]
                #print(nc)
                for caracter in caracteres:
                    r=cadena.find(caracter)
                    if (r!=a):
                        nueva_cadena=cadena.replace(',','.')
                        #print(nueva_cadena)
                        tabla.cell(i,j).text=nueva_cadena

def agregar_parrafo(string,doc):
    p=doc.add_paragraph(string)
    return p

def pdf_a_word(pdf):
    pdfn=pdf.replace('.pdf','')
    doc = aw.Document(pdf)
    doc.save(f'{pdfn}.docx')

def leerpdf(pdf,palabra):
    pdf_documento = pdf
    documento=fitz.open(pdf_documento)
    p=documento.page_count
    a=-1
    c=0
    for i in range(p):
        pagina=documento.load_page(i)
        text=pagina.get_text('text')
        r=text.find(palabra)
        if r!=a:
            #print('Se encontro en la linea '+str(i+1))
            #print('Se encontro')
            c=c+1
    if (c==0):
        print('No se encontro')
    else:
        print('Se encontro '+str(c)+ ' veces')

def extraerTexto_imagen(palabra,imagen):
    image_path =imagen
    path_to_tesseract = r'C:/Program Files (x86)/Tesseract-OCR/tesseract.exe'
    img = Image.open(image_path) 
    pytesseract.tesseract_cmd = path_to_tesseract 
    text = pytesseract.image_to_string(img) 
    r=text.find(palabra) 
    if r!=-1:
        print('Se encontro')
    else:
        print('No encontro')

def pdf_imagen(pdf):
    # import module
    pages = convert_from_path(pdf)
    for i in range(len(pages)):
        pages[i].save('page'+ str(i) +'.jpg', 'JPEG')

def extraerTexto_imagenV2(nombre,cantidad,palabra,tipodoc,expediente):#modificar la ruta de la imagen
    for i in range(cantidad):
        imagen='C:/Users/DELL/Desktop/pruebaword/'+expediente+'/'+tipodoc+'/'+nombre+f'{i+1}.jpg'
        print('Pagina '+str(i+1)) 
        extraerTexto_imagen(palabra,imagen)

